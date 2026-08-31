"""Render a Review's analysis as a portable file — PDF or DOCX.

Owner directive 2026-08-31 ("Export Report … PDF, DOCX") closes 49.12's open
formats question. Two properties are structural here:

* **One content model, two renderers.** ``build_export_model`` assembles the
  document once, from payloads already serialized under the caller's own
  permissions (``serialize_finding`` with the caller's ``legal_position`` flag,
  ``report_payload``). A field the API would omit for this caller — an internal
  legal position above all (LEGAL-02) — is therefore absent from the file too,
  and neither renderer can reintroduce it.
* **Counts, never a grade.** The model carries the same deliberate absences as
  the report endpoint: no risk score, no verdict, no percentage presented as a
  conclusion (36.10, F-8, F-9), and the F-9 sentence travels with the numbers.

Both renderers use libraries already in the ingestion stack (pymupdf,
python-docx) — no new dependency (rule 19).
"""

from __future__ import annotations

import html
import io
from dataclasses import dataclass
from typing import Any

# Fixed rendering order for classification counts — attention first. A rendering
# order, not a severity model: Tier-1 states are legally equivalent (DESIGN.md).
_COUNT_ORDER = (
    "DEVIATION", "MISSING", "CONFLICT", "UNABLE_TO_EVALUATE",
    "AMBIGUOUS", "UNRESOLVED", "MATCH",
)


@dataclass(frozen=True)
class Block:
    """One content block. ``kind``: h1 | h2 | h3 | p | kv | quote | note."""
    kind: str
    text: str
    label: str = ""


def _value_text(value: Any) -> str:
    """A stored JSON value as plain text — the server's keys and values verbatim,
    never interpreted (rule 12)."""
    if value is None:
        return "Not recorded"
    if isinstance(value, dict):
        return "; ".join(f"{k}: {_value_text(v)}" for k, v in value.items())
    if isinstance(value, list):
        return ", ".join(_value_text(v) for v in value)
    return str(value)


def _evidence_location(e: dict[str, Any]) -> str:
    parts: list[str] = []
    if e.get("section_number"):
        parts.append(f"§{e['section_number']}")
    if e.get("section_title"):
        parts.append(str(e["section_title"]))
    if e.get("page_number") is not None:
        parts.append(f"p.{e['page_number']}")
    return " · ".join(parts) if parts else "location not recorded"


def build_export_model(*, contract: dict[str, Any], version: dict[str, Any],
                       review: dict[str, Any], report: dict[str, Any],
                       findings: list[dict[str, Any]],
                       exported_at: str) -> list[Block]:
    blocks: list[Block] = [
        Block("h1", contract.get("name") or "Analysis report"),
        Block("kv", str(contract.get("contract_type") or "not declared"),
              label="Document type"),
        Block("kv", f"v{version.get('version_number')}"
                    f" — {version.get('original_filename') or ''}".rstrip(" —"),
              label="Version"),
        Block("kv", (review.get("completed_at") or review.get("created_at")
                     or "")[:10], label="Analysis date"),
        Block("kv", str(review.get("configuration_snapshot_id") or ""),
              label="Configuration snapshot"),
        Block("kv", str(review.get("status") or ""), label="Review status"),
        Block("kv", exported_at[:16].replace("T", " "), label="Exported"),
        Block("note",
              "This report never grades the document. Counts and findings are "
              "reported exactly as the analysis recorded them; whether a "
              "deviation is acceptable is a Legal Decision made by an "
              "authorized person, never by this report."),
        Block("h2", "Summary"),
    ]

    counts = report.get("classification_counts") or {}
    for classification in _COUNT_ORDER:
        n = counts.get(classification)
        if n:
            blocks.append(Block("kv", str(n), label=classification))
    for classification, n in sorted(counts.items()):
        if classification not in _COUNT_ORDER and n:
            blocks.append(Block("kv", str(n), label=classification))

    coverage = report.get("coverage") or {}
    blocks.append(Block(
        "kv",
        f"{coverage.get('requirements_with_findings', 0)} of "
        f"{coverage.get('requirements_in_snapshot', 0)} requirements in the "
        "snapshot produced findings",
        label="Coverage"))
    blocks.append(Block("kv", str(report.get("findings_requiring_decision", 0)),
                        label="Findings awaiting a Legal Decision"))
    blocks.append(Block("kv", str(report.get("unmatched_provisions", 0)),
                        label="Unmatched provisions (document-level observations)"))

    blocks.append(Block("h2", "Findings"))
    if not findings:
        blocks.append(Block("p", "This Review produced no Findings."))

    for finding in findings:
        requirement = finding.get("requirement") or {}
        title = " — ".join(
            part for part in (requirement.get("code"), requirement.get("name"))
            if part) or "Requirement"
        blocks.append(Block("h3", title))
        blocks.append(Block("kv", str(finding.get("classification")),
                            label="Classification"))
        blocks.append(Block("kv", str(finding.get("status")), label="Status"))
        if finding.get("escalated"):
            blocks.append(Block("kv", "Yes", label="Escalated"))

        for evaluation in finding.get("evaluations") or []:
            scope = evaluation.get("scope_key") or ""
            if evaluation.get("scope_label"):
                scope = f"{scope} · {evaluation['scope_label']}"
            blocks.append(Block("p", f"Evaluation: {scope}"))
            # Presence-tested, exactly like the screens (52.4): a field the API
            # omitted for this caller renders nothing — no placeholder.
            if "rule_outcome" in evaluation:
                blocks.append(Block("kv", str(evaluation["rule_outcome"]),
                                    label="Rule outcome"))
            blocks.append(Block("kv", _value_text(evaluation.get("actual_value")),
                                label="Found in contract"))
            if "expected_value" in evaluation:
                blocks.append(Block("kv", _value_text(evaluation["expected_value"]),
                                    label="Company Standard"))
            if "operator" in evaluation and evaluation.get("operator"):
                blocks.append(Block("kv", str(evaluation["operator"]),
                                    label="Comparison"))
            for line in evaluation.get("explanation") or []:
                blocks.append(Block("p", f"• {_value_text(line)}"))
            decision = evaluation.get("current_decision")
            if decision:
                blocks.append(Block(
                    "kv",
                    f"{decision.get('decision_type')} "
                    f"(v{decision.get('version_number')}, "
                    f"{(decision.get('created_at') or '')[:10]})",
                    label="Legal Decision"))

        evidence = finding.get("evidence") or []
        if evidence:
            blocks.append(Block("p", "Evidence:"))
            for e in evidence:
                blocks.append(Block("quote", str(e.get("content") or ""),
                                    label=_evidence_location(e)))
        else:
            blocks.append(Block(
                "p", "No supporting text was found in the document for this "
                     "Requirement."))

    blocks.append(Block(
        "note",
        "Generated by LegalMind. Contents reflect what the exporting account "
        "is authorized to see; fields withheld from that account are absent "
        "from this file."))
    return blocks


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------
def render_docx(blocks: list[Block]) -> bytes:
    from docx import Document

    doc = Document()
    for block in blocks:
        if block.kind == "h1":
            doc.add_heading(block.text, level=0)
        elif block.kind == "h2":
            doc.add_heading(block.text, level=1)
        elif block.kind == "h3":
            doc.add_heading(block.text, level=2)
        elif block.kind == "kv":
            paragraph = doc.add_paragraph()
            paragraph.add_run(f"{block.label}: ").bold = True
            paragraph.add_run(block.text)
        elif block.kind == "quote":
            if block.label:
                doc.add_paragraph(block.label).runs[0].italic = True
            doc.add_paragraph(block.text, style="Intense Quote")
        elif block.kind == "note":
            doc.add_paragraph(block.text).runs[0].italic = True
        else:
            doc.add_paragraph(block.text)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# --------------------------------------------------------------------------
# PDF — pymupdf Story: HTML assembled from escaped text only.
# --------------------------------------------------------------------------
_PDF_CSS = """
body { font-family: sans-serif; font-size: 10pt; color: #111; }
h1 { font-size: 18pt; margin: 0 0 6pt; }
h2 { font-size: 13pt; margin: 14pt 0 4pt; }
h3 { font-size: 11pt; margin: 10pt 0 3pt; }
p { margin: 2pt 0; }
blockquote { margin: 2pt 0 6pt 12pt; font-style: italic; color: #333; }
.note { font-style: italic; color: #444; margin: 6pt 0; }
.loc { color: #555; font-size: 8.5pt; margin: 4pt 0 0; }
"""


def _blocks_to_html(blocks: list[Block]) -> str:
    parts: list[str] = []
    for block in blocks:
        text = html.escape(block.text)
        label = html.escape(block.label)
        if block.kind in ("h1", "h2", "h3"):
            parts.append(f"<{block.kind}>{text}</{block.kind}>")
        elif block.kind == "kv":
            parts.append(f"<p><b>{label}:</b> {text}</p>")
        elif block.kind == "quote":
            if label:
                parts.append(f'<p class="loc">{label}</p>')
            parts.append(f"<blockquote>{text}</blockquote>")
        elif block.kind == "note":
            parts.append(f'<p class="note">{text}</p>')
        else:
            parts.append(f"<p>{text}</p>")
    return "<body>" + "".join(parts) + "</body>"


def render_pdf(blocks: list[Block]) -> bytes:
    import pymupdf

    story = pymupdf.Story(html=_blocks_to_html(blocks), user_css=_PDF_CSS)
    buffer = io.BytesIO()
    writer = pymupdf.DocumentWriter(buffer)
    page_rect = pymupdf.paper_rect("a4")
    # Rect + tuple is pymupdf's inset arithmetic, not sequence concatenation.
    content_rect = page_rect + (36, 40, -36, -40)  # noqa: RUF005
    more = True
    while more:
        device = writer.begin_page(page_rect)
        more, _ = story.place(content_rect)
        story.draw(device)
        writer.end_page()
    writer.close()
    return buffer.getvalue()


RENDERERS = {"pdf": render_pdf, "docx": render_docx}
MEDIA_TYPES = {
    "pdf": "application/pdf",
    "docx": ("application/vnd.openxmlformats-officedocument"
             ".wordprocessingml.document"),
}
