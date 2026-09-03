"""Document parsing and normalization — locked Step 34.

The governing rule is locked 34.9: **extraction failures never result in
invented text or legal conclusions.** Every branch here either produces text
that genuinely came from the document, or reports failure.

Locked rules implemented:
  34.6   native PDF/DOCX text extraction preferred
  34.7   OCR used when a supported PDF has no usable text
  34.8   OCR-derived content explicitly identified
  34.9   failures never invent text
  34.10  partial extraction explicitly represented
  34.11  pages, sections, paragraphs, tables preserved where available
  34.12  existing clause numbering preserved
  34.13  source locations retained for Evidence
  34.14  original extracted text preserved alongside normalized text
"""

from __future__ import annotations

import re
import shutil
from dataclasses import dataclass, field

from legalmind.domain.enums import EvidenceSourceType, ExtractionStatus
from legalmind.ingestion.validation import DOCX_MIME, PDF_MIME

# A page yielding fewer than this many characters of native text is treated as
# having no usable text (34.7). Deliberately low: the threshold decides whether
# to attempt OCR, never what the text says.
MIN_USABLE_CHARS_PER_PAGE = 20

# --------------------------------------------------------------------------
# Legibility — locked 34.3, "detect when normal extraction is INSUFFICIENT"
# --------------------------------------------------------------------------
# The defect this closes (2026-09-03). `MIN_USABLE_CHARS_PER_PAGE` asks only
# whether text is PRESENT. A PDF whose embedded fonts carry a wrong `/ToUnicode`
# CMap yields plenty of characters that are simply the wrong ones — the glyph
# codes reinterpreted as text. Observed on a real upload: an MSA extracted as
#
#     "MaVWeU SeUYLceV AgUeePeQW ... a cRmSan\ incRUSRUaWed XndeU Whe"
#
# for "Master Services Agreement ... a company incorporated under the". Every
# character at or above 0x6D had been shifted by −29 by the producer's broken
# CMap. That passed the presence check with ~1,960 characters on page 1, so OCR
# was never attempted, and the mojibake flowed into evidence, the clause list,
# the chunk index and — worst — a Review that reported MATCH findings against
# text nobody could read.
#
# A `/ToUnicode` CMap that is present and syntactically valid but semantically
# wrong is indistinguishable from a correct one by inspecting the PDF, so the
# signal has to come from the extracted text itself.
#
# WHY A REPAIR IS NOT ATTEMPTED. Inverting the observed shift was tried as a
# diagnostic and does recover the body text — while corrupting every capital,
# because the shifted range collides with the upper-case block ("Strad" became
# "ptrad"). A remap that silently damages some characters to fix others is
# exactly the invented text 34.9 forbids. OCR reads the RENDERED glyphs, which
# are correct, so it is the only honest route — and it is the route 34.3 already
# prescribes for insufficient extraction.
#
# THE SIGNAL. The share of alphabetic tokens that are common English function
# words. Legal prose is saturated with them; a glyph-code stream is not,
# because the mapping destroys them ("the" → "Whe", "of" → "Rf", "to" → "WR").
# It is a fixed 24-word list, not a dictionary and not a model: deterministic,
# offline, and cheap.
#
# THE THRESHOLD, measured rather than guessed. Across the 21 real PDFs in the
# supplied corpus and the statute set, document-level shares ran 0.210 (a
# policy page dense with product nouns) to 0.422 (Companies Act). The garbled
# upload measured 0.092, and 0.345 once the shift was undone — i.e. inside the
# healthy band, confirming both the signal and the gap. 0.15 sits between the
# two populations: 29% below the lowest legitimate document, 63% above the
# garbled one.
LEGIBILITY_STOPWORDS = frozenset({
    "the", "of", "to", "and", "in", "a", "is", "that", "for", "it", "as",
    "with", "be", "on", "by", "or", "this", "are", "from", "at", "not",
    "which", "shall", "any",
})

#: Below this share of function words, a Latin-script document's extracted text
#: is treated as illegible rather than as content.
ILLEGIBLE_STOPWORD_SHARE = 0.15

#: Judged at DOCUMENT level and never per page. Measured on the same corpus: a
#: cover page, a signature page and a website footer legitimately score 0.055 to
#: 0.086 because they are noun lists with almost no prose. Deciding per page
#: would send correctly-extracted pages to OCR and make good text worse.
MIN_WORDS_TO_JUDGE_LEGIBILITY = 200

#: The share of a text's characters that must be plain ASCII before an English
#: function-word test means anything. A document in another script scores near
#: zero on that list for reasons that have nothing to do with extraction
#: quality, so it is reported as unjudgeable and left exactly as extracted.
MIN_ASCII_SHARE_TO_JUDGE = 0.85


@dataclass
class Segment:
    """One unit of extracted content, destined for document_evidence.

    ``content`` is the normalized text; ``original_content`` is what the parser
    actually returned (34.14 — both are preserved). ``section_number`` is
    captured only when the document states it (34.12 — numbering is preserved,
    never generated).
    """

    content: str
    original_content: str
    source_type: EvidenceSourceType
    page_number: int | None = None
    section_number: str | None = None
    section_title: str | None = None
    start_offset: int | None = None
    end_offset: int | None = None
    metadata: dict = field(default_factory=dict)


@dataclass
class ParseResult:
    segments: list[Segment]
    status: ExtractionStatus
    pages_total: int = 0
    pages_extracted: int = 0
    pages_failed: list[int] = field(default_factory=list)
    diagnostics: list[str] = field(default_factory=list)
    #: True when the caller asked for OCR to be deferred (``defer_ocr``) and this
    #: document needs it. The result then carries NO segments and its ``status``
    #: is provisional — the caller must not persist it as the document's
    #: extraction outcome; the outcome belongs to the later OCR run (42.5 gives
    #: OCR its own ProcessingRunType for exactly this attempt-history shape).
    needs_ocr: bool = False
    #: Where page numbers came from, recorded on the processing run (34.13's
    #: spirit — a location is only as good as its provenance). ``None`` when the
    #: document yielded no page model. PDF pages are physical; DOCX pages come
    #: from the file's OWN pagination record (see _docx_paragraph_pages).
    pagination_source: str | None = None


class ParseError(Exception):
    """Raised only for conditions that make the document unreadable."""


# --------------------------------------------------------------------------
# Normalization — locked 34.12
# --------------------------------------------------------------------------
def normalize_text(raw: str) -> str:
    """Whitespace normalization only.

    Deliberately conservative: it collapses runs of whitespace and normalizes
    line endings. It does NOT correct spelling, expand abbreviations, repair
    OCR errors or alter numbers — locked 45C.18 permits normalizing an OCR
    error only when deterministic, and this layer cannot establish that.
    """
    text = raw.replace("\r\n", "\n").replace("\r", "\n")
    # The third character in this class is U+00A0 NO-BREAK SPACE, deliberately: PDF
    # extraction emits them freely, and a clause whose spacing differs only by an
    # nbsp must normalize to the same text, or `ENG-11` determinism would depend on
    # which producer wrote the file. Not an error to "fix" to an ASCII space.
    text = re.sub(r"[ \t ]+", " ", text)  # noqa: RUF001
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# Clause numbering as documents actually write it: "8.", "8.2", "12.3.4",
# "Section 8", "ARTICLE IV", "(a)". Recognition only — never invention.
_CLAUSE_PATTERNS = (
    re.compile(r"^(?P<num>\d+(?:\.\d+)*)\.?\s+(?P<title>[A-Z][^\n]{0,120})?"),
    re.compile(r"^(?:Section|SECTION|Clause|CLAUSE)\s+(?P<num>\d+(?:\.\d+)*)"
               r"\.?\s*(?P<title>[^\n]{0,120})?"),
    re.compile(r"^(?:Article|ARTICLE)\s+(?P<num>[IVXLC]+|\d+)"
               r"\.?\s*(?P<title>[^\n]{0,120})?"),
)


def detect_clause_number(line: str) -> tuple[str | None, str | None]:
    """Return (section_number, section_title) if the line states one.

    Locked 34.12 — existing clause numbering is *preserved*. If the document
    does not state a number, this returns ``None``; a number is never generated,
    because a fabricated section reference would corrupt evidence traceability.
    """
    stripped = line.strip()
    if not stripped or len(stripped) > 200:
        return None, None
    for pattern in _CLAUSE_PATTERNS:
        m = pattern.match(stripped)
        if m:
            title = (m.group("title") or "").strip(" .:-") or None
            return m.group("num"), title
    return None, None


def segment_paragraphs(text: str, *, page_number: int | None,
                       source_type: EvidenceSourceType,
                       base_offset: int = 0) -> list[Segment]:
    """Split into paragraph segments, carrying source locations (34.13)."""
    segments: list[Segment] = []
    cursor = 0
    for block in re.split(r"\n\s*\n", text):
        raw = block
        start = text.find(raw, cursor)
        if start < 0:                          # pragma: no cover - defensive
            start = cursor
        cursor = start + len(raw)
        normalized = normalize_text(raw)
        if not normalized:
            continue
        first_line = normalized.split("\n", 1)[0]
        number, title = detect_clause_number(first_line)
        segments.append(Segment(
            content=normalized,
            original_content=raw,
            source_type=source_type,
            page_number=page_number,
            section_number=number,
            section_title=title,
            start_offset=base_offset + start,
            end_offset=base_offset + cursor,
        ))
    return segments


# --------------------------------------------------------------------------
# OCR availability — locked 34.7, 34.9
# --------------------------------------------------------------------------
def ocr_available() -> bool:
    """Whether the locked OCR toolchain (OCRmyPDF + Tesseract) is present."""
    return shutil.which("ocrmypdf") is not None and shutil.which("tesseract") is not None


def stopword_share(text: str) -> float | None:
    """Share of alphabetic tokens that are common English function words.

    ``None`` means "not judgeable", for one of two honest reasons: too little
    text to measure, or a text that is not predominantly ASCII (see
    ``MIN_ASCII_SHARE_TO_JUDGE``). A caller must treat ``None`` as "leave the
    extraction alone", never as a failure.
    """
    if not text:
        return None
    printable = [c for c in text if not c.isspace()]
    if not printable:
        return None
    ascii_share = sum(c.isascii() for c in printable) / len(printable)
    if ascii_share < MIN_ASCII_SHARE_TO_JUDGE:
        return None

    words = re.findall(r"[A-Za-z]{1,12}", text.lower())
    if len(words) < MIN_WORDS_TO_JUDGE_LEGIBILITY:
        return None
    return sum(word in LEGIBILITY_STOPWORDS for word in words) / len(words)


def text_is_legible(text: str) -> bool | None:
    """Whether extracted text reads as language rather than as glyph codes.

    ``None`` when the question cannot be answered (see ``stopword_share``), which
    is deliberately distinct from ``False``: only ``False`` is evidence of a
    broken extraction, and only ``False`` may change what the parser does.
    """
    share = stopword_share(text)
    if share is None:
        return None
    return share >= ILLEGIBLE_STOPWORD_SHARE


# --------------------------------------------------------------------------
# Parsers
# --------------------------------------------------------------------------
def parse_pdf(data: bytes, *, defer_ocr: bool = False) -> ParseResult:
    """Native PDF text extraction, with OCR only where a page has none.

    Locked 34.7: OCR is used when a supported PDF does not contain usable text.
    Locked 34.9: if a page has no native text and OCR is unavailable or fails,
    that page is reported as failed — its text is never guessed.

    ``defer_ocr`` (2026-09-03, the ~63s-upload fix): when True and this document
    turns out to need OCR — a page with no usable native text, or a native layer
    that is not legible — and the toolchain is present, the OCR is NOT run here.
    The result comes back with ``needs_ocr=True`` and no segments, so the caller
    can finish the upload quickly and run the OCR as its own background
    processing run (``ProcessingRunType.OCR``, locked 42.5). When the toolchain
    is absent there is nothing to defer to, and the fail-closed behaviour is
    exactly as before.
    """
    import pymupdf

    diagnostics: list[str] = []
    try:
        doc = pymupdf.open(stream=data, filetype="pdf")
    except Exception as exc:
        raise ParseError(f"PDF could not be opened: {type(exc).__name__}") from exc

    if doc.needs_pass:
        # 34.4 lists password-protected among the failure conditions.
        raise ParseError("PDF is password protected")

    segments: list[Segment] = []
    failed_pages: list[int] = []
    deferred_pages: list[int] = []
    extracted_pages = 0
    offset = 0

    # `doc` is a PyMuPDF Document; the library ships no type information, so the
    # page objects are untyped here rather than wrongly typed.
    for index, page in enumerate(doc, start=1):    # type: ignore[arg-type,var-annotated]
        try:
            raw = page.get_text("text") or ""
        except Exception as exc:                # pragma: no cover - defensive
            raw = ""
            diagnostics.append(f"page {index}: extraction error {type(exc).__name__}")

        if len(raw.strip()) >= MIN_USABLE_CHARS_PER_PAGE:
            page_segments = segment_paragraphs(
                raw, page_number=index,
                source_type=EvidenceSourceType.NATIVE_TEXT, base_offset=offset)
            segments.extend(page_segments)
            extracted_pages += 1
            offset += len(raw)
            continue

        # No usable native text on this page (34.7).
        if not ocr_available():
            failed_pages.append(index)
            diagnostics.append(
                f"page {index}: no usable native text and OCR toolchain "
                "unavailable; page not extracted")
            continue

        if defer_ocr:
            # The page needs OCR and the caller asked for that to happen in a
            # later run instead of inline. Recorded, not attempted.
            deferred_pages.append(index)
            continue

        try:
            ocr_text = _ocr_page(page)
        except Exception as exc:
            failed_pages.append(index)
            diagnostics.append(f"page {index}: OCR failed ({type(exc).__name__})")
            continue

        if len(ocr_text.strip()) < MIN_USABLE_CHARS_PER_PAGE:
            failed_pages.append(index)
            diagnostics.append(f"page {index}: OCR produced no usable text")
            continue

        segments.extend(segment_paragraphs(
            ocr_text, page_number=index,
            source_type=EvidenceSourceType.OCR,     # 34.8 — explicitly identified
            base_offset=offset))
        extracted_pages += 1
        offset += len(ocr_text)

    pages_total = doc.page_count

    # Locked 34.3 — "detect when normal extraction is INSUFFICIENT and use OCR
    # where supported". Presence was checked per page above; legibility is
    # checked here, once, over the whole document (see the notes on
    # MIN_WORDS_TO_JUDGE_LEGIBILITY for why never per page).
    native_legible = text_is_legible(
        "\n".join(s.content for s in segments
                  if s.source_type is EvidenceSourceType.NATIVE_TEXT))

    if defer_ocr and ocr_available() and (deferred_pages or native_legible is False):
        # OCR is needed and the toolchain is present — hand the decision back to
        # the caller. Without the toolchain there is nothing to defer to, and the
        # existing fail-closed path below concludes now rather than later.
        if deferred_pages:
            diagnostics.append(
                f"{len(deferred_pages)} page(s) have no usable native text; "
                "OCR deferred to a background run")
        if native_legible is False:
            share = stopword_share("\n".join(
                s.content for s in segments
                if s.source_type is EvidenceSourceType.NATIVE_TEXT))
            diagnostics.append(
                "native text is not legible (function-word share "
                f"{share:.3f} < {ILLEGIBLE_STOPWORD_SHARE}); OCR deferred to a "
                "background run")
        pages_total = doc.page_count
        doc.close()
        return ParseResult(
            segments=[], status=ExtractionStatus.FAILED,  # provisional — see needs_ocr
            pages_total=pages_total, pages_extracted=0, pages_failed=[],
            diagnostics=diagnostics, needs_ocr=True)

    if native_legible is False:
        segments, extracted_pages, failed_pages, offset = _reextract_illegible(
            doc, data, segments, diagnostics)

    doc.close()

    status = _status_for(pages_total, extracted_pages)
    return ParseResult(segments=segments, status=status, pages_total=pages_total,
                       pages_extracted=extracted_pages, pages_failed=failed_pages,
                       diagnostics=diagnostics,
                       pagination_source=PDF_PHYSICAL_PAGES if extracted_pages else None)


def _reextract_illegible(
    doc, data: bytes, native_segments: list[Segment], diagnostics: list[str],
) -> tuple[list[Segment], int, list[int], int]:
    """Second pass for a document whose native text is not language.

    Two outcomes, and never a third:

    * **OCR available and measurably better** — the OCR pass replaces the native
      segments wholesale and is marked ``EvidenceSourceType.OCR`` (34.8, and
      34.3's "never silently treat OCR output as equivalent to clean native
      text"). "Measurably better" is required, not assumed: OCR output that is
      itself illegible is discarded, so this can never trade working text for
      worse text.

    * **Otherwise** — every page is reported FAILED and no text is returned.
      That is 34.9 read strictly: the alternative is passing a glyph-code stream
      into evidence, the clause list, the retrieval index and the evaluator,
      where it becomes MATCH findings against text nobody can read. An empty
      extraction is honest; that is not. ``_status_for`` turns it into
      ``FAILED``, which 45B.7 already routes to ``UNABLE_TO_EVALUATE``.

    Returns ``(segments, pages_extracted, pages_failed, offset)``.
    """
    share = stopword_share("\n".join(s.content for s in native_segments))
    diagnostics.append(
        "native text is not legible (function-word share "
        f"{share:.3f} < {ILLEGIBLE_STOPWORD_SHARE}); the embedded fonts most "
        "likely carry an incorrect ToUnicode mapping")

    all_pages = list(range(1, doc.page_count + 1))
    if not ocr_available():
        diagnostics.append(
            "OCR toolchain unavailable, so no legible text can be produced; "
            "no text is returned rather than returning unreadable text")
        return [], 0, all_pages, 0

    ocr_segments: list[Segment] = []
    ocr_failed: list[int] = []
    ocr_pages = 0
    offset = 0
    # Pages OCR independently, so they run in parallel and are reassembled in
    # page order — same input, same output, same order as the sequential loop
    # this replaces (measured byte-identical on the real 30-page document;
    # 65.6s -> 16.2s on 4 cores). Each worker opens its own document handle from
    # `data` because PyMuPDF objects are not safe to share across threads.
    for index, outcome in _ocr_pages_parallel(data, doc.page_count):
        if isinstance(outcome, Exception):
            ocr_failed.append(index)
            diagnostics.append(f"page {index}: OCR failed ({type(outcome).__name__})")
            continue
        ocr_text = outcome
        if len(ocr_text.strip()) < MIN_USABLE_CHARS_PER_PAGE:
            ocr_failed.append(index)
            diagnostics.append(f"page {index}: OCR produced no usable text")
            continue
        ocr_segments.extend(segment_paragraphs(
            ocr_text, page_number=index,
            source_type=EvidenceSourceType.OCR, base_offset=offset))
        ocr_pages += 1
        offset += len(ocr_text)

    ocr_legible = text_is_legible("\n".join(s.content for s in ocr_segments))
    if ocr_legible is not True:
        # OCR ran and is no better. Keeping the native mojibake would be the
        # worse of two bad options: it looks like content.
        diagnostics.append(
            "OCR did not produce legible text either; no text is returned")
        return [], 0, all_pages, 0

    diagnostics.append(
        f"re-extracted {ocr_pages} of {doc.page_count} page(s) by OCR "
        "(marked OCR-derived); the native text was discarded as unreadable")
    return ocr_segments, ocr_pages, ocr_failed, offset


#: Page-level OCR parallelism. Bounded by the machine, never unbounded: each
#: worker is one tesseract process pinned to one thread (OMP_THREAD_LIMIT=1),
#: so the pool as a whole uses about the same CPU a single unpinned tesseract
#: (~2.5 threads) already did — it just keeps all cores busy for the whole run.
OCR_MAX_WORKERS = 4

#: The render resolution for OCR input. 300dpi is tesseract's recommended
#: input density; measured on the real 30-page document, 150dpi is ~35% faster
#: but produces slightly different output, so the density stays at 300 and the
#: speed comes from parallelism, which changes nothing about the output.
OCR_RENDER_DPI = 300


def _ocr_page(page) -> str:                     # pragma: no cover - needs toolchain
    """OCR a single page via the locked toolchain.

    Only reached when the toolchain is present; absence is handled by the caller
    as a failure, never as empty text.
    """
    pix = page.get_pixmap(dpi=OCR_RENDER_DPI)
    return _tesseract_pixmap(pix)


def _tesseract_pixmap(pix) -> str:              # pragma: no cover - needs toolchain
    import os
    import subprocess
    import tempfile

    # One thread per tesseract process: the parallelism lives at the page level,
    # and an unpinned tesseract would oversubscribe the cores it shares with the
    # other pages' workers.
    env = {**os.environ, "OMP_THREAD_LIMIT": "1"}
    with tempfile.TemporaryDirectory() as tmp:
        image_path = f"{tmp}/page.png"
        pix.save(image_path)
        proc = subprocess.run(
            ["tesseract", image_path, "stdout"],
            capture_output=True, text=True, timeout=120, check=False, env=env)
        if proc.returncode != 0:
            raise ParseError(f"tesseract exited {proc.returncode}")
        return proc.stdout


def _ocr_pages_parallel(
    data: bytes, page_count: int,
) -> list[tuple[int, str | Exception]]:
    """OCR every page of ``data``, up to ``OCR_MAX_WORKERS`` at a time.

    Returns ``[(page_number, text-or-exception), ...]`` in page order regardless
    of completion order, so the caller's output is identical to a sequential
    pass. A page's failure is returned as its exception, never raised — one bad
    page must not cost the other twenty-nine (34.10).
    """
    import os
    from concurrent.futures import ThreadPoolExecutor

    import pymupdf

    def one(index: int) -> str | Exception:     # pragma: no cover - needs toolchain
        try:
            # A thread-local document handle: PyMuPDF documents are not
            # thread-safe, and opening from bytes costs well under a millisecond.
            doc = pymupdf.open(stream=data, filetype="pdf")
            try:
                return _ocr_page(doc[index - 1])
            finally:
                doc.close()
        except Exception as exc:
            return exc

    workers = max(1, min(OCR_MAX_WORKERS, os.cpu_count() or 1))
    if workers == 1 or page_count == 1:
        return [(i, one(i)) for i in range(1, page_count + 1)]
    with ThreadPoolExecutor(max_workers=workers) as pool:
        return list(zip(range(1, page_count + 1),
                        pool.map(one, range(1, page_count + 1)), strict=True))


_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# The values written to the processing run's `pagination_source`.
DOCX_RENDERED_PAGE_BREAKS = "DOCX_RENDERED_PAGE_BREAKS"
DOCX_EXPLICIT_PAGE_BREAKS = "DOCX_EXPLICIT_PAGE_BREAKS"
PDF_PHYSICAL_PAGES = "PDF_PHYSICAL_PAGES"


def _docx_paragraph_pages(paragraphs) -> tuple[list[int] | None, int, str | None]:
    """The starting page of each paragraph, from the document's OWN record.

    A DOCX has no physical pages — pagination happens at render time — so this
    never computes one. It reads two things the file itself carries:

    * ``w:lastRenderedPageBreak`` — where the authoring application (Word)
      recorded that a page boundary fell when the file was last saved. This is
      the same pagination the author saw, including the effect of fonts,
      margins and hard breaks.
    * ``w:br w:type="page"`` — a page break the author explicitly inserted.

    Word records a rendered marker at the boundary a hard break causes, so when
    rendered markers exist, a hard break paired with one (no text between them)
    is ONE boundary, not two — measured on the live MSA: 26 rendered + 6
    explicit − 5 such pairs = 27 boundaries = 28 pages, matching Word. The pair
    can appear in EITHER order — Word records the rendered marker at the point
    the previous page's content ended, which is not always after the break run
    — so pairing is order-independent: whichever of the two comes first opens
    the boundary, and the other (if it follows before any real text) closes the
    same one rather than opening a second.

    A file carrying NEITHER kind of break yields ``(None, 0, None)``: a
    one-page document and a converter that strips pagination metadata (Google
    Docs exports do) are indistinguishable, and locked 34.9/34.12 forbid
    guessing — the viewer then says "Not paginated", which stays true.

    A paragraph is assigned the page it STARTS on: a boundary marker that
    precedes the paragraph's first text means the paragraph begins the new
    page; a boundary after text means it began on the earlier page.
    """
    rendered = explicit = 0
    for para in paragraphs:
        for el in para._p.iter():
            if el.tag == _W + "lastRenderedPageBreak":
                rendered += 1
            elif el.tag == _W + "br" and el.get(_W + "type") == "page":
                explicit += 1
    if rendered == 0 and explicit == 0:
        return None, 0, None

    use_rendered = rendered > 0
    page = 1
    # Which kind of break is "open" — awaiting either real text (which closes
    # it) or the OTHER kind of break (which pairs with it as the same
    # boundary, in whichever order the two appear). Two breaks of the SAME
    # kind in a row, with no text between, are two distinct boundaries — a
    # deliberately blank page — so only cross-kind pairing collapses.
    pending_explicit = False
    pending_rendered = False
    pages: list[int] = []
    for para in paragraphs:
        start_page: int | None = None
        for el in para._p.iter():
            if el.tag == _W + "t":
                if (el.text or "").strip():
                    if start_page is None:
                        start_page = page
                    pending_explicit = False
                    pending_rendered = False
            elif el.tag == _W + "br" and el.get(_W + "type") == "page":
                if pending_rendered:
                    pending_rendered = False  # the same boundary, already counted
                else:
                    page += 1
                pending_explicit = True
            elif use_rendered and el.tag == _W + "lastRenderedPageBreak":
                if pending_explicit:
                    pending_explicit = False  # the same boundary, already counted
                else:
                    page += 1
                pending_rendered = True
        pages.append(start_page if start_page is not None else page)

    source = DOCX_RENDERED_PAGE_BREAKS if use_rendered else DOCX_EXPLICIT_PAGE_BREAKS
    return pages, page, source


def parse_docx(data: bytes) -> ParseResult:
    """DOCX extraction preserving paragraphs and tables (34.11).

    Page numbers (2026-09-02): read from the document's own pagination record —
    see ``_docx_paragraph_pages``. Never computed, never guessed; absent record
    means absent pages. Table segments carry no page: python-docx surfaces
    tables outside the paragraph stream, so a table's position in that record
    is not stated by the file.
    """
    import io

    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ParseError(f"DOCX could not be opened: {type(exc).__name__}") from exc

    paragraphs = list(document.paragraphs)
    paragraph_pages, total_pages, pagination_source = _docx_paragraph_pages(paragraphs)

    segments: list[Segment] = []
    offset = 0

    for index, para in enumerate(paragraphs):
        raw = para.text or ""
        normalized = normalize_text(raw)
        if not normalized:
            offset += len(raw)
            continue
        number, title = detect_clause_number(normalized.split("\n", 1)[0])
        segments.append(Segment(
            content=normalized, original_content=raw,
            source_type=EvidenceSourceType.NATIVE_TEXT,
            page_number=paragraph_pages[index] if paragraph_pages else None,
            section_number=number, section_title=title,
            start_offset=offset, end_offset=offset + len(raw),
            metadata={"style": para.style.name if para.style else None},
        ))
        offset += len(raw)

    # Tables are preserved as TABLE-sourced segments (34.11).
    for table_index, table in enumerate(document.tables, start=1):
        rows = ["\t".join(cell.text.strip() for cell in row.cells)
                for row in table.rows]
        raw = "\n".join(rows)
        normalized = normalize_text(raw)
        if not normalized:
            continue
        segments.append(Segment(
            content=normalized, original_content=raw,
            source_type=EvidenceSourceType.TABLE,
            start_offset=offset, end_offset=offset + len(raw),
            metadata={"table_index": table_index, "rows": len(table.rows)},
        ))
        offset += len(raw)

    status = (ExtractionStatus.COMPLETE if segments else ExtractionStatus.FAILED)
    diagnostics = [] if segments else ["DOCX contained no extractable text"]
    if segments and pagination_source:
        diagnostics = [*diagnostics,
                       f"pages from the document's own pagination record ({pagination_source})"]
    return ParseResult(segments=segments, status=status,
                       pages_total=total_pages if pagination_source else 0,
                       pages_extracted=(total_pages if pagination_source else 1) if segments else 0,
                       diagnostics=diagnostics,
                       pagination_source=pagination_source if segments else None)


def parse(data: bytes, mime_type: str, *, defer_ocr: bool = False) -> ParseResult:
    if mime_type == PDF_MIME:
        return parse_pdf(data, defer_ocr=defer_ocr)
    if mime_type == DOCX_MIME:
        return parse_docx(data)  # a DOCX is text-native; OCR never applies
    raise ParseError(f"unsupported mime type: {mime_type}")


def _status_for(pages_total: int, pages_extracted: int) -> ExtractionStatus:
    """Locked 34.10 — partial extraction is explicitly represented."""
    if pages_total == 0 or pages_extracted == 0:
        return ExtractionStatus.FAILED
    if pages_extracted < pages_total:
        return ExtractionStatus.PARTIAL
    return ExtractionStatus.COMPLETE
