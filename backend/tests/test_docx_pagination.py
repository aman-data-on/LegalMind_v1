"""DOCX pagination from the document's own record — 2026-09-02, owner request.

The rule under test (locked 34.9/34.12 applied to pages): a page number is
REPORTED only when the file itself states where its pages fall — Word's
``w:lastRenderedPageBreak`` markers, or author-inserted ``w:br type="page"``
breaks. Nothing is computed, rendered or guessed; a DOCX carrying neither kind
of break keeps ``page_number = None`` exactly as before, and the viewer keeps
saying "Not paginated".

Fixture DOCX files are built in memory — no real document enters the
repository (locked 54.6).
"""

from __future__ import annotations

import io

import docx
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement

from legalmind.ingestion.parsing import (
    DOCX_EXPLICIT_PAGE_BREAKS,
    DOCX_RENDERED_PAGE_BREAKS,
    parse_docx,
)


def _mark_rendered_break(paragraph, *, before_text: str | None = None) -> None:
    """Inject ``w:lastRenderedPageBreak`` the way Word writes it: inside a run,
    before the text that begins the new page."""
    run = paragraph.add_run(before_text or "")
    marker = OxmlElement("w:lastRenderedPageBreak")
    run._r.insert(0, marker)


def _build(paragraph_builder) -> bytes:
    d = docx.Document()
    paragraph_builder(d)
    buffer = io.BytesIO()
    d.save(buffer)
    return buffer.getvalue()


def test_a_docx_with_no_break_record_keeps_pages_none() -> None:
    data = _build(lambda d: [d.add_paragraph(f"Paragraph {i} with enough words to matter.")
                             for i in range(4)])
    result = parse_docx(data)
    assert result.pagination_source is None
    assert result.pages_total == 0
    assert all(s.page_number is None for s in result.segments)


def test_rendered_markers_paginate_and_a_marker_leading_a_paragraph_starts_the_new_page() -> None:
    def build(d):
        d.add_paragraph("Title on page one.")
        d.add_paragraph("Still page one.")
        p = d.add_paragraph("")           # begins page two: marker precedes its text
        _mark_rendered_break(p, before_text="1. DEFINITIONS AND INTERPRETATION")
        d.add_paragraph("Body on page two.")

    result = parse_docx(_build(build))
    assert result.pagination_source == DOCX_RENDERED_PAGE_BREAKS
    assert result.pages_total == 2
    pages = [s.page_number for s in result.segments]
    assert pages == [1, 1, 2, 2]
    # The boundary paragraph itself is on the NEW page — the off-by-one that
    # would put "1. DEFINITIONS" on page 1 is the exact bug this test pins.
    heading = next(s for s in result.segments if "DEFINITIONS" in s.content)
    assert heading.page_number == 2


def test_a_break_after_text_leaves_that_paragraph_on_its_starting_page() -> None:
    def build(d):
        p = d.add_paragraph("This paragraph starts on page one")
        _mark_rendered_break(p, before_text=" and spills onto page two.")
        d.add_paragraph("Fully on page two.")

    result = parse_docx(_build(build))
    pages = [s.page_number for s in result.segments]
    assert pages == [1, 2]  # the spilling paragraph keeps its STARTING page


def test_a_hard_break_followed_by_words_rendered_marker_is_one_boundary_not_two() -> None:
    def build(d):
        d.add_paragraph("Page one.")
        d.add_page_break()                 # hard break (its own empty paragraph)
        p = d.add_paragraph("")
        _mark_rendered_break(p, before_text="Page two starts here.")  # Word's record of the SAME boundary
        d.add_paragraph("Also page two.")

    result = parse_docx(_build(build))
    assert result.pages_total == 2         # 3 would mean the boundary was double-counted
    texted = [s.page_number for s in result.segments]
    assert texted == [1, 2, 2]


def test_a_rendered_marker_preceding_its_hard_break_is_still_one_boundary() -> None:
    """The pair can appear in EITHER order. Word can record the rendered marker
    at the point the previous page's content ended, which is not always after
    the break run — so this is the mirror of the "hard break followed by
    rendered marker" case above, and must count the same: one boundary, not
    two. (Regression: an earlier version of this dedup only handled the
    explicit-then-rendered order and double-counted this one, skipping a page
    number in the output.)"""
    def build(d):
        d.add_paragraph("Page one text.")
        p = d.add_paragraph("Tail of page one.")
        _mark_rendered_break(p, before_text="")
        run = p.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)
        d.add_paragraph("Page two text.")

    result = parse_docx(_build(build))
    assert result.pages_total == 2
    assert [s.page_number for s in result.segments] == [1, 1, 2]


def test_explicit_breaks_alone_paginate_when_no_rendered_record_exists() -> None:
    def build(d):
        d.add_paragraph("Page one.")
        d.add_page_break()
        d.add_paragraph("Page two.")
        d.add_page_break()
        d.add_paragraph("Page three.")

    result = parse_docx(_build(build))
    assert result.pagination_source == DOCX_EXPLICIT_PAGE_BREAKS
    assert result.pages_total == 3
    assert [s.page_number for s in result.segments] == [1, 2, 3]


def test_tables_state_no_page_of_their_own() -> None:
    def build(d):
        d.add_paragraph("Page one.")
        d.add_page_break()
        d.add_paragraph("Page two.")
        table = d.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Fee"
        table.rows[0].cells[1].text = "12 months"

    result = parse_docx(_build(build))
    table_segments = [s for s in result.segments if s.metadata.get("table_index")]
    assert table_segments, "the table must still be extracted"
    # python-docx surfaces tables outside the paragraph stream, so the file's
    # break record does not state the table's page — None, never a guess.
    assert all(s.page_number is None for s in table_segments)


def test_upload_ceiling_default_is_25mb(monkeypatch) -> None:
    monkeypatch.delenv("LEGALMIND_MAX_UPLOAD_BYTES", raising=False)
    from legalmind.config import max_upload_bytes
    assert max_upload_bytes() == 25 * 1024 * 1024


def test_qn_helper_matches_the_namespace_the_parser_walks() -> None:
    """A tripwire: the parser matches raw Clark-notation tags; if python-docx's
    namespace ever changed, both this and the parser would need the new URI."""
    from legalmind.ingestion.parsing import _W
    assert qn("w:lastRenderedPageBreak") == _W + "lastRenderedPageBreak"
