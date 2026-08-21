"""Document ingestion tests — locked Step 34.

The rule under most of these is locked 34.9: **extraction failures never result
in invented text or legal conclusions.**
"""

from __future__ import annotations

import io
import zipfile

import pytest

from legalmind.db import models as M
from legalmind.domain import enums as E
from legalmind.ingestion import parsing
from legalmind.ingestion.service import (
    ingest_document,
    process_document_version,
)
from legalmind.ingestion.storage import LocalFilesystemStorage, fingerprint
from legalmind.ingestion.validation import (
    DOCX_MIME,
    PDF_MIME,
    UploadRejected,
    validate_upload,
)
from tests.conftest import make_user


@pytest.fixture
def storage(tmp_path):
    return LocalFilesystemStorage(tmp_path / "objects")


@pytest.fixture
def contract(db):
    owner = make_user(db)
    c = M.Contract(owner_id=owner.id, name="ACME MSA",
                   status=E.ContractStatus.ACTIVE)
    db.add(c); db.flush()
    c._owner = owner
    return c


# ----------------------------------------------------------------- builders
def build_pdf(pages: list[str]) -> bytes:
    import pymupdf
    doc = pymupdf.open()
    for text in pages:
        page = doc.new_page()
        page.insert_text((72, 96), text, fontsize=11)
    data = doc.tobytes()
    doc.close()
    return data


def build_docx(paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    import docx
    d = docx.Document()
    for p in paragraphs:
        d.add_paragraph(p)
    if table:
        t = d.add_table(rows=len(table), cols=len(table[0]))
        for r, row in enumerate(table):
            for c, val in enumerate(row):
                t.cell(r, c).text = val
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def build_image_only_pdf() -> bytes:
    """A PDF page with no text layer — the scanned-document case (34.3)."""
    import pymupdf
    doc = pymupdf.open()
    doc.new_page()          # blank page, no text
    data = doc.tobytes()
    doc.close()
    return data


# ================================================================ validation
def test_rejects_unsupported_type(db):
    with pytest.raises(UploadRejected) as e:
        validate_upload(b"plain text", "notes.txt", "text/plain")
    assert e.value.code == "UNSUPPORTED_TYPE"


def test_rejects_content_type_mismatch(db):
    """34.16 — the client's declared type is never trusted."""
    docx_bytes = build_docx(["hello"])
    with pytest.raises(UploadRejected) as e:
        validate_upload(docx_bytes, "trick.pdf", PDF_MIME)
    assert e.value.code == "CONTENT_TYPE_MISMATCH"


def test_rejects_zip_that_is_not_docx(db):
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as z:
        z.writestr("evil.sh", "#!/bin/sh\n")
    with pytest.raises(UploadRejected) as e:
        validate_upload(buf.getvalue(), "x.docx", DOCX_MIME)
    assert e.value.code == "UNRECOGNISED_CONTENT"


def test_rejects_empty_file(db):
    with pytest.raises(UploadRejected) as e:
        validate_upload(b"", "empty.pdf", PDF_MIME)
    assert e.value.code == "EMPTY_FILE"


# =================================================================== storage
def test_original_is_preserved_byte_identical(storage):
    """34.5 / 34.18 — the ingestion layer must not alter the original."""
    data = build_pdf(["8.1 Term. This Agreement commences on the Effective Date."])
    key = storage.put(data, suggested_name="msa.pdf")
    assert storage.get(key) == data
    assert fingerprint(storage.get(key)) == fingerprint(data)


def test_storage_is_write_once(storage):
    """No update or overwrite operation exists (34.18)."""
    assert not hasattr(storage, "update")
    assert not hasattr(storage, "delete")
    data = b"%PDF-1.7 a"
    k1 = storage.put(data, suggested_name="a.pdf")
    k2 = storage.put(data, suggested_name="a.pdf")
    assert k1 != k2                 # identical content never collides
    assert storage.get(k1) == storage.get(k2) == data


# ================================================================= ingestion
def test_ingest_pdf_produces_evidence_with_locations(db, storage, contract):
    """34.11 / 34.13 — pages and source locations retained for Evidence."""
    data = build_pdf([
        "8.1 Limitation of Liability\n\nThe aggregate liability shall not exceed "
        "six months of fees paid under this Agreement.",
        "12.1 Governing Law\n\nThis Agreement is governed by the laws of India.",
    ])
    result = ingest_document(
        db, storage, contract_id=contract.id, uploaded_by=contract._owner.id,
        data=data, filename="msa.pdf", declared_mime=PDF_MIME)

    assert result.document_version.extraction_status is E.ExtractionStatus.COMPLETE
    assert result.document_version.processing_status is E.ProcessingStatus.COMPLETED
    assert result.evidence_count > 0

    evidence = db.query(M.DocumentEvidence).filter_by(
        processing_run_id=result.processing_run.id).all()
    assert {e.page_number for e in evidence} == {1, 2}
    assert all(e.source_type is E.EvidenceSourceType.NATIVE_TEXT for e in evidence)
    assert all(e.start_offset is not None for e in evidence)


def test_clause_numbering_is_preserved_not_invented(db, storage, contract):
    """34.12 — existing numbering preserved; nothing generated where absent.

    Uses DOCX because it has a real paragraph model; PDF text layout gives no
    reliable paragraph boundaries, so asserting paragraph-level segmentation
    against a synthetic PDF would test the fixture rather than the parser.
    """
    data = build_docx([
        "8.2 Limitation of Liability",
        "A paragraph with no clause number at all.",
        "Section 14 Governing Law",
    ])
    result = ingest_document(
        db, storage, contract_id=contract.id, uploaded_by=contract._owner.id,
        data=data, filename="msa.docx", declared_mime=DOCX_MIME)

    evidence = db.query(M.DocumentEvidence).filter_by(
        processing_run_id=result.processing_run.id).all()
    numbers = {e.section_number for e in evidence}
    assert "8.2" in numbers
    assert "14" in numbers
    assert None in numbers          # unnumbered text stays unnumbered


def test_original_text_preserved_alongside_normalized(db, storage, contract):
    """34.14 — both representations are kept."""
    data = build_docx(["8.1   Term.    Spacing    is   irregular."])
    result = ingest_document(
        db, storage, contract_id=contract.id, uploaded_by=contract._owner.id,
        data=data, filename="msa.docx", declared_mime=DOCX_MIME)
    e = db.query(M.DocumentEvidence).filter_by(
        processing_run_id=result.processing_run.id).first()
    assert "  " not in e.content                                    # normalized
    assert "   " in e.evidence_metadata["original_content"]         # original kept


def test_docx_tables_preserved_as_table_evidence(db, storage, contract):
    """34.11 — tables preserved where technically available."""
    data = build_docx(["1. Fees"], table=[["Item", "Amount"], ["Licence", "1000"]])
    result = ingest_document(
        db, storage, contract_id=contract.id, uploaded_by=contract._owner.id,
        data=data, filename="fees.docx", declared_mime=DOCX_MIME)
    kinds = {e.source_type for e in db.query(M.DocumentEvidence).filter_by(
        processing_run_id=result.processing_run.id).all()}
    assert E.EvidenceSourceType.TABLE in kinds


def test_version_numbers_are_system_controlled_and_sequential(db, storage, contract):
    """42.4 UNIQUE(contract_id, version_number); Step 33.6."""
    for expected in (1, 2, 3):
        r = ingest_document(
            db, storage, contract_id=contract.id, uploaded_by=contract._owner.id,
            data=build_pdf([f"Version {expected} text content here."]),
            filename="msa.pdf", declared_mime=PDF_MIME)
        assert r.document_version.version_number == expected


def test_duplicate_is_detected_not_silently_suppressed(db, storage, contract):
    """34.5 — duplicates are detectable. Whether a re-upload is a new
    contractual version is a business decision (Step 33.9), so ingestion
    reports the duplicate rather than deciding."""
    data = build_pdf(["8.1 Liability capped at six months of fees."])
    first = ingest_document(
        db, storage, contract_id=contract.id, uploaded_by=contract._owner.id,
        data=data, filename="msa.pdf", declared_mime=PDF_MIME)
    second = ingest_document(
        db, storage, contract_id=contract.id, uploaded_by=contract._owner.id,
        data=data, filename="msa-again.pdf", declared_mime=PDF_MIME)

    assert second.duplicate_of == first.document_version.id
    assert second.document_version.id != first.document_version.id
    assert second.document_version.file_hash == first.document_version.file_hash


def test_same_file_in_two_contracts_is_not_a_duplicate(db, storage, contract):
    """42.4 — file_hash is indexed but NOT globally unique: the same source file
    may legitimately appear in multiple contracts."""
    other = M.Contract(owner_id=contract._owner.id, name="Other MSA",
                       status=E.ContractStatus.ACTIVE)
    db.add(other); db.flush()
    data = build_pdf(["8.1 Liability capped at six months of fees."])

    ingest_document(db, storage, contract_id=contract.id,
                    uploaded_by=contract._owner.id, data=data,
                    filename="a.pdf", declared_mime=PDF_MIME)
    r2 = ingest_document(db, storage, contract_id=other.id,
                         uploaded_by=contract._owner.id, data=data,
                         filename="a.pdf", declared_mime=PDF_MIME)
    assert r2.duplicate_of is None


# ============================================== fail-closed extraction (34.9)
def test_no_text_layer_and_no_ocr_fails_closed(db, storage, contract):
    """34.4 / 34.7 / 34.9 — the scanned-document case.

    With no native text and no OCR toolchain, extraction FAILS. It must not
    produce empty-but-successful evidence, because a Requirement evaluated
    against silently-missing text would yield a false MISSING rather than
    UNABLE_TO_EVALUATE (34.17).
    """
    assert parsing.ocr_available() is False      # documents this environment
    result = ingest_document(
        db, storage, contract_id=contract.id, uploaded_by=contract._owner.id,
        data=build_image_only_pdf(), filename="scan.pdf", declared_mime=PDF_MIME)

    assert result.document_version.extraction_status is E.ExtractionStatus.FAILED
    assert result.document_version.processing_status is E.ProcessingStatus.FAILED
    assert result.processing_run.status is E.ProcessingRunStatus.FAILED
    assert result.processing_run.error_code == "EXTRACTION_FAILED"
    assert result.evidence_count == 0            # nothing invented
    assert any("OCR" in d for d in result.diagnostics)


def test_partial_extraction_is_explicitly_represented(db, storage, contract):
    """34.10 — a document where only some pages yield text."""
    import pymupdf
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 96), "8.1 Liability is capped at six months of fees.",
                     fontsize=11)
    doc.new_page()                               # second page: no text layer
    data = doc.tobytes()
    doc.close()

    result = ingest_document(
        db, storage, contract_id=contract.id, uploaded_by=contract._owner.id,
        data=data, filename="mixed.pdf", declared_mime=PDF_MIME)

    assert result.document_version.extraction_status is E.ExtractionStatus.PARTIAL
    assert result.processing_run.run_metadata["pages_failed"] == [2]
    assert result.evidence_count > 0             # page 1 still usable


def test_corrupt_pdf_fails_without_inventing_text(db, storage, contract):
    result = ingest_document(
        db, storage, contract_id=contract.id, uploaded_by=contract._owner.id,
        data=b"%PDF-1.7\nthis is not really a pdf body at all",
        filename="broken.pdf", declared_mime=PDF_MIME)
    assert result.document_version.extraction_status is E.ExtractionStatus.FAILED
    assert result.evidence_count == 0


def test_retry_creates_a_new_run_and_preserves_history(db, storage, contract):
    """42.5 — Attempt 1 FAILED, Attempt 2 COMPLETED must both remain visible."""
    result = ingest_document(
        db, storage, contract_id=contract.id, uploaded_by=contract._owner.id,
        data=build_image_only_pdf(), filename="scan.pdf", declared_mime=PDF_MIME)
    first_run = result.processing_run
    assert first_run.status is E.ProcessingRunStatus.FAILED

    second_run = process_document_version(
        db, storage, result.document_version,
        run_type=E.ProcessingRunType.REPROCESS)

    runs = db.query(M.DocumentProcessingRun).filter_by(
        document_version_id=result.document_version.id).all()
    assert len(runs) == 2
    assert second_run.id != first_run.id
    assert first_run.status is E.ProcessingRunStatus.FAILED   # history intact


def test_extraction_status_is_separate_from_review_lifecycle(db, storage, contract):
    """34.15 / Step 30 r13 — a document-level failure is not a Finding, and
    ANALYSIS_FAILED is not UNABLE_TO_EVALUATE."""
    result = ingest_document(
        db, storage, contract_id=contract.id, uploaded_by=contract._owner.id,
        data=build_image_only_pdf(), filename="scan.pdf", declared_mime=PDF_MIME)
    dv = result.document_version
    assert isinstance(dv.extraction_status, E.ExtractionStatus)
    # The document carries no review status and no finding classification.
    assert not hasattr(dv, "status")
    assert not hasattr(dv, "classification")


def test_ocr_derived_content_would_be_labelled(db):
    """34.8 — OCR-derived content is explicitly identified.

    The toolchain is absent here, so this asserts the labelling contract rather
    than running OCR: OCR segments carry source_type OCR, which is a distinct
    enum member from NATIVE_TEXT.
    """
    from legalmind.ingestion.parsing import Segment
    seg = Segment(content="x", original_content="x",
                  source_type=E.EvidenceSourceType.OCR)
    assert seg.source_type is not E.EvidenceSourceType.NATIVE_TEXT
    assert E.EvidenceSourceType.OCR.value == "OCR"


def test_normalization_never_alters_numbers_or_words(db):
    """45C.18 — normalizing an OCR error is permitted only when deterministic;
    this layer cannot establish that, so it only touches whitespace."""
    raw = "Liability shall not exceed 6 m0nths of fees"
    assert parsing.normalize_text(raw) == raw     # '6 m0nths' left untouched
